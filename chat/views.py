from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth import login, logout
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.conf import settings
from django.db import transaction
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import io
import json
import joblib
import re
import os
import pandas as pd
import uuid
from pathlib import Path
from .models import ChatSession, Message, UploadedFile
from .utils.embeddings import search_knowledge_base, generate_answer_gemini, create_embedding
from .utils.pdf_processor import process_pdf, create_pdf_embeddings, chunk_text
from .utils.video_processor import process_video, create_video_embeddings
from .forms import CustomUserCreationForm

def landing(request):
    """Landing page - redirects to main if authenticated"""
    if request.user.is_authenticated:
        return redirect('rag_main')
    return render(request, 'landing.html')

def register(request):
    """User registration"""
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('rag_main')
    else:
        form = CustomUserCreationForm()
    return render(request, 'registration/register.html', {'form': form})

@login_required
def rag_main(request):
    """Main RAG Dashboard"""
    return render(request, 'chat/rag_main.html')

@login_required
def logout_view(request):
    """Logout and redirect to landing"""
    logout(request)
    return redirect('landing')

def index(request):
    """Render the main chat interface"""
    return render(request, 'chat/index.html')

try:
    from .utils.txt_processor import process_txt
    TXT_AVAILABLE = True
except ImportError:
    TXT_AVAILABLE = False
    print("⚠️ TXT processor not available")

try:
    from .utils.docx_processor import process_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ DOCX processor not available")

try:
    from .utils.pptx_processor import process_pptx
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False
    print("⚠️ PPTX processor not available")

MAX_FILE_SIZE = {
    'video': 5 * 1024 * 1024 * 1024,  
    'pdf': 500 * 1024 * 1024,         
    'image': 100 * 1024 * 1024,        
    'docx': 200 * 1024 * 1024,        
    'txt': 500 * 1024 * 1024,         
    'default': 500 * 1024 * 1024     
}

ALLOWED_EXTENSIONS = {
    'pdf': ['.pdf'],
    'video': ['.mp4', '.avi', '.mov', '.mkv', '.webm'],
    'audio': ['.mp3', '.wav', '.m4a', '.aac', '.flac', '.ogg'],
    'docx': ['.docx', '.doc'],
    'txt': ['.txt'],
    'pptx': ['.pptx', '.ppt'],
}

def validate_file_upload(uploaded_file):
    """Validate uploaded file size and type"""
    if not uploaded_file:
        return None, "No file provided"
    
    ext = Path(uploaded_file.name).suffix.lower()
    file_type = None
    
    for ftype, extensions in ALLOWED_EXTENSIONS.items():
        if ext in extensions:
            file_type = ftype
            break
    
    if not file_type:
        return None, f"Unsupported file type: {ext}"
    
    if file_type == 'audio':
        file_type = 'video'
    if file_type == 'pptx':
        pass 
    
    max_size = MAX_FILE_SIZE.get(file_type, MAX_FILE_SIZE['default'])
    
    if uploaded_file.size > max_size:
        size_mb = uploaded_file.size / (1024 * 1024)
        limit_mb = max_size / (1024 * 1024)
        
        if size_mb > 1024:
            size_gb = size_mb / 1024
            limit_gb = limit_mb / 1024
            return None, f"File too large ({size_gb:.2f}GB). Maximum size for {file_type} is {limit_gb:.1f}GB"
        else:
            return None, f"File too large ({size_mb:.1f}MB). Maximum size for {file_type} is {limit_mb:.0f}MB"
    
    return file_type, None


def load_knowledge_base():
    """Load all embeddings from storage"""
    embeddings_path = settings.DATA_DIR / 'embeddings.joblib'
    pdf_embeddings_path = settings.DATA_DIR / 'pdf_embeddings.joblib'
    
    all_dfs = []
    
    if embeddings_path.exists():
        try:
            video_df = joblib.load(embeddings_path)
            print(f"📚 Loaded video embeddings: {len(video_df)} chunks")
            all_dfs.append(video_df)
        except Exception as e:
            print(f"⚠️ Error loading video embeddings: {e}")
    
    if pdf_embeddings_path.exists():
        try:
            pdf_df = joblib.load(pdf_embeddings_path)
            print(f"📚 Loaded PDF embeddings: {len(pdf_df)} chunks")
            all_dfs.append(pdf_df)
        except Exception as e:
            print(f"⚠️ Error loading PDF embeddings: {e}")
    
    if not all_dfs:
        return None
    
    if len(all_dfs) > 1:
        combined_df = pd.concat(all_dfs, ignore_index=True)
        print(f"📊 Combined knowledge base: {len(combined_df)} total chunks")
        return combined_df
    
    return all_dfs[0]


def save_embeddings(chunks, file_type):
    """Save embeddings to appropriate file"""
    if file_type in ['video']:
        embeddings_file = 'embeddings.joblib'
    else:
        embeddings_file = 'pdf_embeddings.joblib'
    
    embeddings_path = settings.DATA_DIR / embeddings_file
    
    if embeddings_path.exists():
        print(f"📚 Loading existing {embeddings_file}...")
        existing_df = joblib.load(embeddings_path)
        new_df = pd.DataFrame(chunks)
        df = pd.concat([existing_df, new_df], ignore_index=True)
        print(f"   - Merged: {len(existing_df)} + {len(new_df)} = {len(df)} chunks")
    else:
        print(f"📚 Creating new {embeddings_file}...")
        df = pd.DataFrame(chunks)
    
    joblib.dump(df, embeddings_path)
    print(f"💾 Saved {len(chunks)} chunks to {embeddings_file}")
    
    return len(chunks)


def generate_mcqs_response(question):
    """Generate MCQs from question"""
    num_match = re.search(r'\d+', question)
    num_questions = int(num_match.group()) if num_match else 5
    num_questions = min(max(num_questions, 1), 20)
    
    df = load_knowledge_base()
    
    context = ""
    if df is not None:
        sample_size = min(5, len(df))
        sample_chunks = df.sample(n=sample_size)
        context = "\n\n".join(sample_chunks['text'].values)
    
    prompt = f"""Generate {num_questions} multiple choice questions based on the following content.

{"Content from knowledge base:" if context else "Topic:"} 
{context if context else question}

Format each question as:
Q1. [Question text]
A) [Option A]
B) [Option B]
C) [Option C]
D) [Option D]
✓ Correct Answer: [Letter]
Explanation: [Brief explanation]

Make questions challenging and educational."""
    
    result = generate_answer_gemini(prompt, "", has_context=False, similarity_score=0.0)
    
    return {
        'answer': result['answer'],
        'sources': [{'title': 'Generated MCQs', 'chunk': f'{num_questions} Questions'}]
    }

def generate_diagram_response(question):
    """Generate ASCII diagram from question"""
    topic = question.lower()
    for keyword in ['diagram of', 'diagram for', 'draw', 'visualize', 'concept map of', 'mind map of', 'flowchart of', 'show me']:
        if keyword in topic:
            topic = topic.split(keyword)[-1].strip()
            break
    
    prompt = f"""Create a detailed ASCII art diagram for: {topic}. Return ONLY the ASCII diagram, no extra explanation."""
    
    result = generate_answer_gemini(prompt, "", has_context=False, similarity_score=0.0)
    
    return {
        'answer': result['answer'],
        'sources': [{'title': 'Generated ASCII Diagram', 'chunk': 'Text-Based'}]
    }

def process_normal_question(question, user):
    """Process questions - only show sources when HIGHLY confident"""
    df = load_knowledge_base()
    
    if df is None:
        result = generate_answer_gemini(question, "", has_context=False, similarity_score=0.0)
        return {'answer': result['answer'], 'sources': []}
    
    if 'user_id' in df.columns:
        df = df[df['user_id'] == user.id]
    
    if len(df) == 0:
        result = generate_answer_gemini(question, "", has_context=False, similarity_score=0.0)
        return {'answer': result['answer'], 'sources': []}
    
    search_results = search_knowledge_base(df, question)
    max_similarity = search_results['max_similarity']
    
    context = "\n\n".join(search_results['results']['text'].values)
    
    result = generate_answer_gemini(
        question, 
        context, 
        has_context=search_results['has_context'],
        similarity_score=max_similarity
    )
    
    sources = []
    
    if search_results['has_context'] and max_similarity >= 0.65:
        print(f"   ✅ KB source (confidence: {max_similarity:.3f})")
        sources = [{
            'title': row['title'],
            'chunk': row['number'],
            'source_type': 'kb'         
        } for _, row in search_results['results'].iterrows()]
    else:
        print(f"   ℹ️ General knowledge answer (confidence: {max_similarity:.3f})")
        sources = [{
            'title': 'Generated Content',
            'chunk': 'AI Generated',
            'source_type': 'generated' 
    }]
    
    return {
        'answer': result['answer'],
        'sources': sources
    }

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def create_session(request):
    """Create a new chat session for logged-in user"""
    try:
        session_id = str(uuid.uuid4())
        session = ChatSession.objects.create(
            session_id=session_id,
            title="New Chat",
            user=request.user
        )
        return JsonResponse({
            'success': True,
            'session_id': session.session_id,
            'title': session.title
        })
    except Exception as e:
        print(f"❌ Error creating session: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@require_http_methods(["GET"])
@login_required
def get_sessions(request):
    """Get list of chat sessions for logged-in user ONLY"""
    try:
        sessions = ChatSession.objects.filter(user=request.user).order_by('-updated_at')[:20]
        return JsonResponse({
            'success': True,
            'sessions': [{
                'id': s.session_id,
                'title': s.title,
                'updated_at': s.updated_at.isoformat()
            } for s in sessions]
        })
    except Exception as e:
        print(f"❌ Error getting sessions: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@require_http_methods(["GET"])
@login_required
def get_messages(request, session_id):
    """Get all messages for a specific session"""
    try:
        session = ChatSession.objects.get(session_id=session_id, user=request.user)
        messages = session.messages.all().order_by('timestamp')
        return JsonResponse({
            'success': True,
            'messages': [{
                'role': m.role,
                'content': m.content,
                'sources': m.sources if m.sources else [],
                'timestamp': m.timestamp.isoformat()
            } for m in messages]
        })
    except ChatSession.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Session not found'}, status=404)
    except Exception as e:
        print(f"❌ Error getting messages: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["DELETE"])
@login_required
def delete_session(request, session_id):
    """Delete a chat session"""
    try:
        session = ChatSession.objects.get(session_id=session_id, user=request.user)
        session.delete()
        return JsonResponse({'success': True})
    except ChatSession.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Session not found'}, status=404)
    except Exception as e:
        print(f"❌ Error deleting session: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def send_message(request):
    """Handle incoming chat messages and generate responses"""
    try:
        data = json.loads(request.body)
        session_id = data.get('session_id')
        question = data.get('message', '').strip()
        
        if not session_id or not question:
            return JsonResponse({
                'success': False,
                'error': 'Missing session_id or message'
            }, status=400)
        
        print(f"\n{'='*60}")
        print(f"📨 New question: {question}")
        print(f"{'='*60}")
        
        session, created = ChatSession.objects.get_or_create(
            session_id=session_id,
            user=request.user,
            defaults={'title': 'New Chat'}
        )
        
        Message.objects.create(
            session=session,
            role='user',
            content=question
        )
        
        answer_text = None
        sources = []
        
        if 'youtube.com' in question.lower() or 'youtu.be' in question.lower():
            print("🎬 YouTube URL detected - processing transcript...")
            
            video_id_match = re.search(r'(?:v=|\/|youtu\.be\/)([0-9A-Za-z_-]{11})', question)
            
            if video_id_match:
                video_id = video_id_match.group(1)
                print(f"   Video ID: {video_id}")
                
                try:
                    from .utils.video_processor import get_youtube_transcript, create_video_embeddings
                    
                    result = get_youtube_transcript(video_id)
                    
                    if result:
                        title = f"YouTube_{video_id}"
                        chunks = result['chunks']
                        
                        for chunk in chunks:
                            chunk['title'] = f"{title} (Video)"
                            chunk['user_id'] = request.user.id
                        
                        chunks_with_embeddings = create_video_embeddings(chunks)
                        chunks_saved = save_embeddings(chunks_with_embeddings, 'video')
                        
                        caption_type = result.get('caption_type', 'unknown')
                        caption_quality = 'HIGH (Manual Captions)' if caption_type == 'manual' else 'MEDIUM (Auto-Generated)'
                        
                        answer_text = f"""✅ **Successfully processed YouTube video!**

📺 **Video ID:** `{video_id}`
🎯 **Transcript Source:** {result.get('source', 'YouTube API')}
📊 **Caption Quality:** {caption_quality}
📦 **Chunks Created:** {chunks_saved}
📝 **Total Content:** {len(result['full_transcript']):,} characters

The video transcript has been added to your knowledge base and is ready for questions!

💡 **Try asking:**
• "What is this video about?"
• "Summarize the main points from the video"
• "Explain [specific topic mentioned in video]"
• "What are the key takeaways?"
"""
                        
                        sources = [{
                            'title': f'YouTube Video {video_id}',
                            'chunk': f'{chunks_saved} chunks ({caption_type} captions)'
                        }]
                        
                        print(f"✅ YouTube video processed successfully!")
                    
                    else:
                        answer_text = """⚠️ **Could not retrieve YouTube transcript**

This usually happens when:
• The video doesn't have captions/subtitles enabled
• Captions are disabled by the video creator
• The video is private, unlisted, or age-restricted
• The video is from a live stream without captions

**What you can do:**
1. Try a different video with captions enabled
2. Check if the video has the "CC" button on YouTube
3. If you own the video, enable auto-captions in YouTube Studio

I can still help answer questions if you describe the video content to me!"""
                        
                        sources = [{'title': 'YouTube Processing', 'chunk': 'Transcript unavailable'}]
                
                except Exception as e:
                    print(f"❌ YouTube processing error: {e}")
                    import traceback
                    traceback.print_exc()
                    
                    answer_text = f"""❌ **Error processing YouTube video**

An unexpected error occurred while processing the video.

Error: {str(e)}

Please try again or try a different video."""
                    sources = [{'title': 'Processing Error', 'chunk': str(e)[:100]}]
            
            else:
                answer_text = """⚠️ **Invalid YouTube URL**

I detected a YouTube link but couldn't extract the video ID.

**Supported formats:**
• `https://www.youtube.com/watch?v=VIDEO_ID`
• `https://youtu.be/VIDEO_ID`
• `https://m.youtube.com/watch?v=VIDEO_ID`

Please make sure you're using a valid YouTube URL."""
                sources = [{'title': 'Invalid URL', 'chunk': 'Could not parse'}]
        
        elif any(keyword in question.lower() for keyword in ['mcq', 'multiple choice', 'quiz', 'test questions', 'practice questions']):
            print("📝 MCQ request detected")
            response = generate_mcqs_response(question)
            answer_text = response['answer']
            sources = response['sources']
        
        elif any(keyword in question.lower() for keyword in ['diagram', 'draw', 'visualize', 'concept map', 'mind map', 'flowchart', 'show me']):
            print("📊 Diagram request detected")
            response = generate_diagram_response(question)
            answer_text = response['answer']
            sources = response['sources']
        
        else:
            print("💬 Processing normal question")
            response = process_normal_question(question, request.user)
            answer_text = response['answer']
            sources = response['sources']
        
        if session.messages.filter(role='user').count() == 1:
            session.title = question[:50] + "..." if len(question) > 50 else question
            session.save()
        
        Message.objects.create(
            session=session,
            role='assistant',
            content=answer_text,
            sources=sources
        )
        
        print(f"✅ Response generated successfully")
        print(f"{'='*60}\n")
        
        return JsonResponse({
            'success': True,
            'answer': answer_text,
            'sources': sources or []
        })
        
    except ChatSession.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Session not found'}, status=404)
    except Exception as e:
        print(f"❌ Error in send_message: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def upload_file(request):
    """Handle file uploads for logged-in user"""
    try:
        uploaded_file = request.FILES.get('file')
        
        file_type, error = validate_file_upload(uploaded_file)
        if error:
            return JsonResponse({'success': False, 'error': error}, status=400)

        if file_type == 'audio':
            file_type = 'video'

        file_obj = UploadedFile.objects.create(
            user=request.user,
            file_type=file_type,
            original_filename=uploaded_file.name,
            file_path=uploaded_file,
            processing_status='uploaded'
        )
        
        print(f"✅ File uploaded by {request.user.username} (ID: {file_obj.id})")
        
        return JsonResponse({
            'success': True,
            'file_id': file_obj.id,
            'filename': file_obj.original_filename,
            'type': file_type,
            'status': 'uploaded'
        })
        
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def process_file(request, file_id):
    """Process uploaded file and create embeddings with file_id tracking"""
    try:
        file_obj = UploadedFile.objects.get(id=file_id, user=request.user)
        file_obj.processing_status = 'processing'
        file_obj.save()
        
        print(f"\n{'='*70}")
        print(f"🔄 Starting file processing: {file_obj.original_filename}")
        print(f"   Type: {file_obj.file_type}")
        print(f"   File ID: {file_obj.id}")
        print(f"{'='*70}")
        
        file_path = file_obj.file_path.path
        chunks_with_embeddings = []
        title = Path(file_obj.original_filename).stem
        
        if file_obj.file_type == 'pdf':
            print("📄 Processing as PDF...")
            result = process_pdf(file_path, file_obj.original_filename)
            chunks_with_embeddings = create_pdf_embeddings(result['chunks'])
            title = result.get('title', title)
            
        elif file_obj.file_type == 'video':
            print("🎬 Processing as video...")
            result = process_video(file_path, file_obj.original_filename)
            chunks_with_embeddings = create_video_embeddings(result['chunks'])
            title = result.get('title', title)
        
        elif file_obj.file_type == 'docx':
            if not DOCX_AVAILABLE:
                raise Exception("DOCX processing not available.")
            
            print("📝 Processing as DOCX...")
            result = process_docx(file_path, file_obj.original_filename)
            
            if not result.get('success'):
                raise Exception(result.get('error', 'DOCX processing failed'))
            
            chunks = chunk_text(result['text'], chunk_size=800, overlap=100)
            for chunk in chunks:
                chunk['title'] = f"{result['title']} (DOCX)"
            
            chunks_with_embeddings = create_pdf_embeddings(chunks)
            title = result['title']

        elif file_obj.file_type == 'txt':
            if not TXT_AVAILABLE:
                raise Exception("TXT processing not available.")
    
            print("📝 Processing as TXT...")
            result = process_txt(file_path, file_obj.original_filename)
    
            if not result.get('chunks'):
                raise Exception("No content could be extracted from TXT file")
    
            chunks = result['chunks']
    
            for chunk in chunks:
                chunk['file_id'] = file_obj.id
                chunk['original_filename'] = file_obj.original_filename
    
            chunks_with_embeddings = create_pdf_embeddings(chunks)
            title = result['title']
            
        elif file_obj.file_type == 'pptx':
            if not PPTX_AVAILABLE:
                raise Exception("PPTX processing not available. Run: pip install python-pptx")

            print("📊 Processing as PPTX...")
            result = process_pptx(file_path, file_obj.original_filename)

            if not result.get('success'):
                raise Exception(result.get('error', 'PPTX processing failed'))

            chunks = result['chunks']

            for chunk in chunks:
                chunk['file_id'] = file_obj.id
                chunk['original_filename'] = file_obj.original_filename

            chunks_with_embeddings = create_pdf_embeddings(chunks)
            title = result['title']

        else:
            raise Exception(f"Unsupported file type: {file_obj.file_type}")
        
        if not chunks_with_embeddings:
            raise Exception("No content could be extracted from file")
        
        print(f"\n📊 Processing summary:")
        print(f"   - Chunks created: {len(chunks_with_embeddings)}")
        print(f"   - Title: {title}")
        print(f"   - File ID: {file_obj.id}")
        
        for chunk in chunks_with_embeddings:
            if 'file_id' not in chunk:
                chunk['file_id'] = file_obj.id
            if 'user_id' not in chunk:
                chunk['user_id'] = file_obj.user.id
            if 'original_filename' not in chunk:
                chunk['original_filename'] = file_obj.original_filename
        
        chunks_saved = save_embeddings(chunks_with_embeddings, file_obj.file_type)
        
        file_obj.processed = True
        file_obj.processing_status = 'completed'
        file_obj.chunks_count = chunks_saved
        file_obj.save()
        
        print(f"✅ File processing complete!")
        print(f"   - Database record updated: ID={file_obj.id}")
        print(f"   - Status: {file_obj.processing_status}")
        print(f"   - Chunks: {file_obj.chunks_count}")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'status': 'completed',
            'chunks': chunks_saved,
            'title': title
        })
        
    except UploadedFile.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'File not found'}, status=404)
    except Exception as e:
        print(f"\n❌ ERROR in process_file:")
        print(f"{'='*70}")
        import traceback
        traceback.print_exc()
        print(f"{'='*70}\n")
        
        if 'file_obj' in locals():
            file_obj.processing_status = f'error: {str(e)}'
            file_obj.save()
        
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["GET"])
@login_required
def get_processing_status(request, file_id):
    """Get real-time processing status"""
    try:
        file_obj = UploadedFile.objects.get(id=file_id, user=request.user)
        
        return JsonResponse({
            'success': True,
            'status': file_obj.processing_status,
            'processed': file_obj.processed,
            'chunks': file_obj.chunks_count or 0
        })
        
    except UploadedFile.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'File not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
    
@require_http_methods(["GET"])
@login_required
def get_knowledge_base_status(request):
    """Get knowledge base statistics for logged-in user ONLY"""
    try:
        files = UploadedFile.objects.filter(user=request.user)
        total_chunks = sum(f.chunks_count or 0 for f in files)
        
        print(f"📊 KB status for {request.user.username}: {files.count()} files, {total_chunks} chunks")
        
        return JsonResponse({
            'success': True,
            'total_files': files.count(),
            'processed_files': files.filter(processed=True).count(),
            'total_chunks': total_chunks,
            'files': [{
                'id': f.id,
                'name': f.original_filename,
                'type': f.file_type,
                'status': f.processing_status,
                'chunks': f.chunks_count or 0
            } for f in files]
        })
    except Exception as e:
        print(f"❌ Error in get_knowledge_base_status: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@require_http_methods(["GET"])
@login_required
def get_knowledge_base_details(request):
    """Get detailed information about logged-in user's files ONLY"""
    try:
        embeddings_path = settings.DATA_DIR / 'embeddings.joblib'
        pdf_embeddings_path = settings.DATA_DIR / 'pdf_embeddings.joblib'
        
        uploaded_files = UploadedFile.objects.filter(user=request.user)
        user_file_ids = set(uploaded_files.values_list('id', flat=True))
        
        print(f"\n{'='*70}")
        print(f"📊 Loading KB Details for {request.user.username}")
        print(f"   User file IDs from DB: {user_file_ids}")
        print(f"{'='*70}")
        
        all_files = {}
        
        if embeddings_path.exists():
            try:
                video_df = joblib.load(embeddings_path)
                print(f"📂 Video embeddings: {len(video_df)} total chunks")
                
                if 'user_id' in video_df.columns:
                    video_df = video_df[video_df['user_id'] == request.user.id]
                    print(f"   After user_id filter: {len(video_df)} chunks")
                elif 'file_id' in video_df.columns:
                    video_df = video_df[video_df['file_id'].isin(user_file_ids)]
                    print(f"   After file_id filter: {len(video_df)} chunks")
                
                for title in video_df['title'].unique():
                    chunks = video_df[video_df['title'] == title]
                    
                    file_obj = None
                    first_chunk = chunks.iloc[0]
                    
                    if 'file_id' in first_chunk and pd.notna(first_chunk['file_id']):
                        try:
                            file_obj = uploaded_files.get(id=int(first_chunk['file_id']))
                        except Exception as e:
                            print(f"   ⚠️ Couldn't find file_id {first_chunk['file_id']}: {e}")
                            continue
                    
                    if not file_obj:
                        print(f"   ⚠️ Skipping '{title}' - no matching file object")
                        continue
                    
                    all_files[title] = {
                        'id': file_obj.id,
                        'type': 'video',
                        'title': title,
                        'total_chunks': len(chunks),
                        'chunks': [
                            {
                                'chunk_id': idx,
                                'number': row['number'],
                                'text': row['text'][:200] + '...' if len(row['text']) > 200 else row['text'],
                                'full_text': row['text'],
                                'start': row.get('start', 'N/A'),
                                'end': row.get('end', 'N/A')
                            }
                            for idx, row in chunks.iterrows()
                        ]
                    }
                    print(f"   ✅ Added: {title} ({len(chunks)} chunks)")
                    
            except Exception as e:
                print(f"⚠️ Error processing video embeddings: {e}")
                import traceback
                traceback.print_exc()
        
        if pdf_embeddings_path.exists():
            try:
                pdf_df = joblib.load(pdf_embeddings_path)
                print(f"\n📂 PDF embeddings: {len(pdf_df)} total chunks")
                
                if 'user_id' in pdf_df.columns:
                    pdf_df = pdf_df[pdf_df['user_id'] == request.user.id]
                    print(f"   After user_id filter: {len(pdf_df)} chunks")
                elif 'file_id' in pdf_df.columns:
                    pdf_df = pdf_df[pdf_df['file_id'].isin(user_file_ids)]
                    print(f"   After file_id filter: {len(pdf_df)} chunks")
                
                for title in pdf_df['title'].unique():
                    chunks = pdf_df[pdf_df['title'] == title]
                    
                    file_obj = None
                    first_chunk = chunks.iloc[0]
                    
                    if 'file_id' in first_chunk and pd.notna(first_chunk['file_id']):
                        try:
                            file_obj = uploaded_files.get(id=int(first_chunk['file_id']))
                        except Exception as e:
                            print(f"   ⚠️ Couldn't find file_id {first_chunk['file_id']}: {e}")
                            continue
                    
                    if not file_obj:
                        print(f"   ⚠️ Skipping '{title}' - no matching file object")
                        continue
                    
                    all_files[title] = {
                        'id': file_obj.id,
                        'type': file_obj.file_type,
                        'title': title,
                        'total_chunks': len(chunks),
                        'chunks': [
                            {
                                'chunk_id': idx,
                                'number': row['number'],
                                'text': row['text'][:200] + '...' if len(row['text']) > 200 else row['text'],
                                'full_text': row['text']
                            }
                            for idx, row in chunks.iterrows()
                        ]
                    }
                    print(f"   ✅ Added: {title} ({len(chunks)} chunks)")
                    
            except Exception as e:
                print(f"⚠️ Error processing PDF embeddings: {e}")
                import traceback
                traceback.print_exc()
        
        print(f"\n{'='*70}")
        print(f"📊 FINAL Knowledge Base Summary for {request.user.username}:")
        print(f"   - Total files: {len(all_files)}")
        print(f"   - Total chunks: {sum(f['total_chunks'] for f in all_files.values())}")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'files': list(all_files.values()),
            'total_files': len(all_files),
            'total_chunks': sum(f['total_chunks'] for f in all_files.values())
        })
        
    except Exception as e:
        print(f"❌ Error getting KB details: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@require_http_methods(["GET"])
@login_required
def get_chunk_details(request, chunk_id):
    """Get full details of a specific chunk"""
    try:
        embeddings_path = settings.DATA_DIR / 'embeddings.joblib'
        pdf_embeddings_path = settings.DATA_DIR / 'pdf_embeddings.joblib'
        
        if embeddings_path.exists():
            try:
                df = joblib.load(embeddings_path)
                if chunk_id in df.index:
                    row = df.loc[chunk_id]
                    return JsonResponse({
                        'success': True,
                        'chunk': {
                            'title': row['title'],
                            'number': row['number'],
                            'text': row['text'],
                            'type': 'video',
                            'start': row.get('start', 'N/A'),
                            'end': row.get('end', 'N/A')
                        }
                    })
            except Exception as e:
                print(f"⚠️ Error searching video embeddings: {e}")
        
        if pdf_embeddings_path.exists():
            try:
                df = joblib.load(pdf_embeddings_path)
                if chunk_id in df.index:
                    row = df.loc[chunk_id]
                    return JsonResponse({
                        'success': True,
                        'chunk': {
                            'title': row['title'],
                            'number': row['number'],
                            'text': row['text'],
                            'type': 'pdf'
                        }
                    })
            except Exception as e:
                print(f"⚠️ Error searching PDF embeddings: {e}")
        
        return JsonResponse({'success': False, 'error': 'Chunk not found'}, status=404)
        
    except Exception as e:
        print(f"❌ Error in get_chunk_details: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["DELETE"])
@transaction.atomic
@login_required
def delete_uploaded_file(request, file_id):
    """Delete an uploaded file completely"""
    try:
        file_obj = UploadedFile.objects.get(id=file_id, user=request.user)
        filename = file_obj.original_filename
        file_type = file_obj.file_type
        
        print(f"\n{'='*70}")
        print(f"🗑️  Deleting file: {filename}")
        print(f"   File ID: {file_id}")
        print(f"   File Type: {file_type}")
        print(f"{'='*70}")
        
        if file_obj.file_path:
            try:
                if hasattr(file_obj.file_path, 'path'):
                    file_path = file_obj.file_path.path
                else:
                    file_path = str(file_obj.file_path)
                
                print(f"   File path: {file_path}")
                
                if os.path.exists(file_path) and not file_path.startswith('http'):
                    os.remove(file_path)
                    print(f"✅ Deleted physical file: {file_path}")
                elif file_path.startswith('http'):
                    print(f"⏭️  Skipping URL (no physical file): {file_path}")
                else:
                    print(f"⚠️ Physical file not found: {file_path}")
                    
            except Exception as e:
                print(f"⚠️ Could not delete physical file: {e}")
                import traceback
                traceback.print_exc()
        
        if file_type == 'video':
            embeddings_file = 'embeddings.joblib'
        else:
            embeddings_file = 'pdf_embeddings.joblib'
        
        embeddings_path = settings.DATA_DIR / embeddings_file
        chunks_removed = 0
        
        if embeddings_path.exists():
            try:
                print(f"\n📂 Loading {embeddings_file}...")
                df = joblib.load(embeddings_path)
                initial_count = len(df)
                print(f"   Initial chunks: {initial_count}")
                
                if 'file_id' in df.columns:
                    print(f"   Using file_id matching (file_id={file_id})...")
                    df_filtered = df[df['file_id'] != file_id]
                    chunks_removed_by_id = initial_count - len(df_filtered)
                    
                    if chunks_removed_by_id > 0:
                        df = df_filtered
                        print(f"   ✅ Removed {chunks_removed_by_id} chunks by file_id")
                        chunks_removed = chunks_removed_by_id
                
                if chunks_removed == 0 and 'original_filename' in df.columns:
                    print(f"   Trying original_filename matching ({filename})...")
                    df_filtered = df[df['original_filename'] != filename]
                    chunks_removed_by_filename = initial_count - len(df_filtered)
                    
                    if chunks_removed_by_filename > 0:
                        df = df_filtered
                        print(f"   ✅ Removed {chunks_removed_by_filename} chunks by filename")
                        chunks_removed = chunks_removed_by_filename
                
                if chunks_removed == 0:
                    print(f"   Trying title pattern matching...")
                    base_title = Path(filename).stem
                    patterns_to_match = [
                        f"{base_title} (Video)",
                        f"{base_title} (PDF)",
                        f"{base_title} (DOCX)",
                        f"{base_title} (TXT)",
                        f"{base_title} (Text)",
                    ]
                    
                    mask = df['title'].isin(patterns_to_match)
                    df_filtered = df[~mask]
                    chunks_removed_by_title = initial_count - len(df_filtered)
                    
                    if chunks_removed_by_title > 0:
                        df = df_filtered
                        print(f"   ✅ Removed {chunks_removed_by_title} chunks by title")
                        chunks_removed = chunks_removed_by_title
                
                final_count = len(df)
                
                if final_count > 0:
                    joblib.dump(df, embeddings_path)
                    print(f"\n💾 Updated {embeddings_file}:")
                    print(f"   Before: {initial_count} chunks")
                    print(f"   After: {final_count} chunks")
                    print(f"   Removed: {chunks_removed} chunks")
                else:
                    embeddings_path.unlink()
                    print(f"\n🗑️  Deleted {embeddings_file} (was empty)")
                    chunks_removed = initial_count
                
            except Exception as e:
                print(f"❌ Error updating embeddings: {e}")
                import traceback
                traceback.print_exc()
        
        file_obj.delete()
        print(f"✅ Deleted from database (ID: {file_id})")
        
        print(f"\n{'='*70}")
        print(f"✅ DELETION COMPLETE")
        print(f"   File: {filename}")
        print(f"   Chunks removed: {chunks_removed}")
        print(f"   Database: ✓")
        print(f"   Embeddings: ✓")
        print(f"   Physical file: ✓")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'message': f'Successfully deleted {filename}',
            'chunks_removed': chunks_removed
        })
        
    except UploadedFile.DoesNotExist:
        print(f"❌ File not found: ID={file_id}")
        return JsonResponse({
            'success': False,
            'error': 'File not found'
        }, status=404)
        
    except Exception as e:
        print(f"\n{'='*70}")
        print(f"❌ ERROR deleting file:")
        print(f"   File ID: {file_id}")
        print(f"   Error: {str(e)}")
        print(f"{'='*70}\n")
        
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'success': False,
            'error': f'Deletion failed: {str(e)}'
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def delete_file_by_title(request):
    """Delete file by title (for legacy files without IDs)"""
    try:
        data = json.loads(request.body)
        title = data.get('title', '').strip()
        
        if not title:
            return JsonResponse({'success': False, 'error': 'Title required'}, status=400)
        
        print(f"\n{'='*70}")
        print(f"🗑️ Deleting by title: {title}")
        print(f"{'='*70}")
        
        if '(Video)' in title:
            embeddings_file = 'embeddings.joblib'
        else:
            embeddings_file = 'pdf_embeddings.joblib'
        
        embeddings_path = settings.DATA_DIR / embeddings_file
        chunks_removed = 0
        
        if embeddings_path.exists():
            df = joblib.load(embeddings_path)
            initial_count = len(df)
            
            df = df[df['title'] != title]
            chunks_removed = initial_count - len(df)
            
            if len(df) > 0:
                joblib.dump(df, embeddings_path)
            else:
                embeddings_path.unlink()
            
            print(f"✅ Removed {chunks_removed} chunks")
        
        try:
            clean_title = title.replace(' (Video)', '').replace(' (PDF)', '').replace(' (DOCX)', '').replace(' (TXT)', '').replace(' (Text)', '')
            UploadedFile.objects.filter(original_filename__icontains=clean_title, user=request.user).delete()
        except:
            pass
        
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'message': f'Deleted {chunks_removed} chunks for "{title}"',
            'chunks_removed': chunks_removed
        })
        
    except Exception as e:
        print(f"❌ Delete by title error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def cleanup_embeddings_api(request):
    """Remove orphaned chunks from embeddings - syncs with database"""
    try:
        print(f"\n{'='*70}")
        print(f"🧹 Starting Embeddings Cleanup (API Call)")
        print(f"{'='*70}")
        
        valid_file_ids = set(UploadedFile.objects.filter(user=request.user).values_list('id', flat=True))
        valid_filenames = set(UploadedFile.objects.filter(user=request.user).values_list('original_filename', flat=True))
        
        print(f"\n📊 Database Status:")
        print(f"   Valid files: {len(valid_file_ids)}")
        
        total_chunks_removed = 0
        files_cleaned = []
        
        embeddings_files = [
            ('embeddings.joblib', 'video'),
            ('pdf_embeddings.joblib', 'documents')
        ]
        
        for embeddings_filename, category in embeddings_files:
            embeddings_path = settings.DATA_DIR / embeddings_filename
            
            if not embeddings_path.exists():
                print(f"⏭️  Skipping {embeddings_filename} (doesn't exist)")
                continue
            
            print(f"\n📂 Processing: {embeddings_filename}")
            
            try:
                df = joblib.load(embeddings_path)
                initial_count = len(df)
                initial_files = df['title'].nunique() if 'title' in df.columns else 0
                
                print(f"   Initial: {initial_count} chunks, {initial_files} files")
                
                if initial_count == 0:
                    print(f"   ✓ Already empty")
                    continue
                
                if 'file_id' in df.columns:
                    print(f"   Using file_id matching...")
                    df_cleaned = df[df['file_id'].isin(valid_file_ids)]
                    chunks_removed = initial_count - len(df_cleaned)
                    
                    if chunks_removed > 0:
                        print(f"   ✅ Removed {chunks_removed} orphaned chunks by file_id")
                        df = df_cleaned
                
                elif 'original_filename' in df.columns:
                    print(f"   Using filename matching...")
                    df_cleaned = df[df['original_filename'].isin(valid_filenames)]
                    chunks_removed = initial_count - len(df_cleaned)
                    
                    if chunks_removed > 0:
                        print(f"   ✅ Removed {chunks_removed} orphaned chunks by filename")
                        df = df_cleaned
                
                else:
                    print(f"   Using title pattern matching...")
                    
                    valid_titles = set()
                    for file_obj in UploadedFile.objects.filter(user=request.user):
                        base_name = Path(file_obj.original_filename).stem
                        
                        if file_obj.file_type == 'video':
                            valid_titles.add(f"{base_name} (Video)")
                        elif file_obj.file_type == 'pdf':
                            valid_titles.add(f"{base_name} (PDF)")
                        elif file_obj.file_type == 'docx':
                            valid_titles.add(f"{base_name} (DOCX)")
                        elif file_obj.file_type == 'txt':
                            valid_titles.add(f"{base_name} (TXT)")
                    
                    if 'title' in df.columns:
                        df_cleaned = df[df['title'].isin(valid_titles)]
                        chunks_removed = initial_count - len(df_cleaned)
                        
                        if chunks_removed > 0:
                            print(f"   ✅ Removed {chunks_removed} orphaned chunks by title")
                            df = df_cleaned
                
                final_count = len(df)
                chunks_removed_total = initial_count - final_count
                total_chunks_removed += chunks_removed_total
                
                if final_count > 0:
                    joblib.dump(df, embeddings_path)
                    final_files = df['title'].nunique() if 'title' in df.columns else 0
                    print(f"   Final: {final_count} chunks, {final_files} files")
                    
                    files_cleaned.append({
                        'file': embeddings_filename,
                        'category': category,
                        'removed': chunks_removed_total,
                        'remaining': final_count
                    })
                else:
                    embeddings_path.unlink()
                    print(f"   🗑️  Deleted {embeddings_filename} (empty)")
                    
                    files_cleaned.append({
                        'file': embeddings_filename,
                        'category': category,
                        'removed': chunks_removed_total,
                        'remaining': 0,
                        'deleted': True
                    })
                
            except Exception as e:
                print(f"   ❌ Error: {e}")
                import traceback
                traceback.print_exc()
        
        print(f"\n{'='*70}")
        print(f"✅ CLEANUP COMPLETE")
        print(f"   Total orphaned chunks removed: {total_chunks_removed}")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'chunks_removed': total_chunks_removed,
            'files_cleaned': files_cleaned,
            'message': f'Successfully removed {total_chunks_removed} orphaned chunks'
        })
        
    except Exception as e:
        print(f"❌ Cleanup error: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def cleanup_media_files(request):
    """Clean up orphaned files in media/uploads"""
    try:
        print(f"\n{'='*70}")
        print(f"🧹 Cleaning up media/uploads folder")
        print(f"{'='*70}")
        
        uploads_dir = settings.MEDIA_ROOT / 'uploads'
        
        if not uploads_dir.exists():
            return JsonResponse({
                'success': True,
                'message': 'No uploads directory found',
                'deleted': 0
            })
        
        db_files = set()

        for file_obj in UploadedFile.objects.all():  
            if file_obj.file_path and hasattr(file_obj.file_path, 'path'):
                try:
                    db_files.add(Path(file_obj.file_path.path).name)
                except:
                    pass
        
        print(f"   Database has {len(db_files)} files")
        
        deleted_count = 0
        deleted_files = []
        
        for file_path in uploads_dir.glob('*'):
            if file_path.is_file():
                filename = file_path.name
                
                if filename in db_files:
                    continue
                
                try:
                    file_path.unlink()
                    deleted_count += 1
                    deleted_files.append(filename)
                    print(f"   ✅ Deleted: {filename}")
                except Exception as e:
                    print(f"   ⚠️ Could not delete {filename}: {e}")
        
        print(f"\n✅ Cleanup complete: {deleted_count} orphaned files removed")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'deleted': deleted_count,
            'files': deleted_files,
            'message': f'Removed {deleted_count} orphaned files from uploads'
        })
        
    except Exception as e:
        print(f"❌ Media cleanup error: {e}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def cleanup_database(request):
    """Delete orphaned database records (files that failed to process)"""
    try:
        print(f"\n{'='*70}")
        print(f"🗑️  Database Cleanup Started")
        print(f"{'='*70}")
        
        unprocessed = UploadedFile.objects.filter(processed=False, user=request.user)
        count = unprocessed.count()
        
        print(f"Found {count} unprocessed files")
        
        deleted_files = []
        
        for file_obj in unprocessed:
            filename = file_obj.original_filename
            print(f"  Deleting: {filename}")
            
            try:
                if file_obj.file_path:
                    file_obj.file_path.delete(save=False)
            except Exception as e:
                print(f"    ⚠️ Could not delete physical file: {e}")
            
            file_obj.delete()
            deleted_files.append(filename)
        
        print(f"✅ Deleted {count} orphaned files")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'deleted': count,
            'files': deleted_files,
            'message': f'Deleted {count} orphaned database records'
        })
        
    except Exception as e:
        print(f"❌ Database cleanup error: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def export_chat_pdf(request, session_id):
    """Export chat session as PDF"""
    try:
        session = ChatSession.objects.get(session_id=session_id, user=request.user)
        messages = session.messages.all().order_by('timestamp')
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        title = Paragraph(f"<b>{session.title}</b>", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        for msg in messages:
            if msg.role == 'user':
                p = Paragraph(f"<b>Question:</b> {msg.content}", styles['Normal'])
            else:
                content = msg.content.replace('<', '&lt;').replace('>', '&gt;')
                p = Paragraph(f"<b>Answer:</b> {content}", styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        
        filename = re.sub(r'[^\w\s-]', '', session.title)[:50]
        
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
        
    except ChatSession.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Session not found'}, status=404)
    except Exception as e:
        print(f"❌ Error exporting PDF: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def export_chat_word(request, session_id):
    """Export chat session as Word document"""
    try:
        session = ChatSession.objects.get(session_id=session_id, user=request.user)
        messages = session.messages.all().order_by('timestamp')
        
        doc = Document()
        
        title = doc.add_heading(session.title, 0)
        
        for msg in messages:
            if msg.role == 'user':
                p = doc.add_paragraph()
                p.add_run('Question: ').bold = True
                p.add_run(msg.content)
            else:
                p = doc.add_paragraph()
                p.add_run('Answer: ').bold = True
                p.add_run(msg.content)
            doc.add_paragraph()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = re.sub(r'[^\w\s-]', '', session.title)[:50]
        
        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response
        
    except ChatSession.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Session not found'}, status=404)
    except Exception as e:
        print(f"❌ Error exporting Word: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def generate_mcqs(request):
    """Generate MCQs from a topic using Gemini"""
    try:
        data = json.loads(request.body)
        topic = data.get('topic', '').strip()
        num_questions = int(data.get('num_questions', 5))
        
        if not topic:
            return JsonResponse({'success': False, 'error': 'Topic is required'}, status=400)
        
        question = f"Generate {num_questions} MCQs about {topic}"
        response = generate_mcqs_response(question)
        
        return JsonResponse({
            'success': True,
            'mcqs': response['answer']
        })
        
    except Exception as e:
        print(f"❌ Error generating MCQs: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def generate_concept_map(request):
    """Generate concept map/mind map using Mermaid"""
    try:
        data = json.loads(request.body)
        topic = data.get('topic', '').strip()
        
        if not topic:
            return JsonResponse({'success': False, 'error': 'Topic is required'}, status=400)
        
        prompt = f"""Create a concept map for "{topic}" in Mermaid.js format.

Use this structure:
```mermaid
graph TD
    A[{topic}] --> B[Concept 1]
    A --> C[Concept 2]
    B --> D[Detail 1]
    B --> E[Detail 2]
```

Make it comprehensive with 10-15 nodes."""
        
        result = generate_answer_gemini(prompt, "", has_context=False, similarity_score=0.0)
        
        return JsonResponse({
            'success': True,
            'diagram': result['answer']
        })
        
    except Exception as e:
        print(f"❌ Error generating concept map: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
@login_required
def download_youtube_video(request):
    """Smart YouTube processor - tries transcript API first, falls back to download"""
    try:
        data = json.loads(request.body)
        url = data.get('url', '').strip()
        
        if not url:
            return JsonResponse({'success': False, 'error': 'URL is required'}, status=400)
        
        if 'youtube.com' not in url.lower() and 'youtu.be' not in url.lower():
            return JsonResponse({
                'success': False,
                'error': 'Please provide a valid YouTube URL'
            }, status=400)
        
        print(f"\n{'='*70}")
        print(f"📺 Processing YouTube Video")
        print(f"   URL: {url}")
        print(f"{'='*70}")
        
        import re
        video_id = None
        patterns = [
            r'(?:v=|\/)([0-9A-Za-z_-]{11})',
            r'^([0-9A-Za-z_-]{11})$'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                video_id = match.group(1)
                if len(video_id) == 11:
                    break
        
        if not video_id:
            return JsonResponse({
                'success': False,
                'error': 'Could not extract video ID from URL'
            }, status=400)
        
        print(f"📺 Video ID: {video_id}")
        
        print(f"\n🚀 STRATEGY 1: Trying YouTube Transcript API...")
        
        try:
            from .utils.video_processor import get_youtube_transcript, create_video_embeddings
            
            result = get_youtube_transcript(video_id)
            
            if result and result.get('chunks'):
                print(f"✅ SUCCESS with YouTube Transcript API!")
                print(f"   - Caption type: {result.get('caption_type', 'unknown')}")
                print(f"   - Chunks: {len(result['chunks'])}")
                
                video_title = result.get('title', f"YouTube_{video_id}")
                clean_title = re.sub(r'[^\w\s-]', '', video_title).strip()
                words = clean_title.split()[:4]  
                clean_title = '_'.join(words)
                title = clean_title if clean_title else f"YouTube_{video_id}"
                
                chunks = result['chunks']
                
                for chunk in chunks:
                    chunk['title'] = f"{title} (Video)"
                    chunk['user_id'] = request.user.id
                
                chunks_with_embeddings = create_video_embeddings(chunks)
                chunks_saved = save_embeddings(chunks_with_embeddings, 'video')
                
                file_obj = UploadedFile.objects.create(
                    user=request.user,
                    file_type='video',
                    original_filename=f"{title}.txt", 
                    file_path=url,
                    processed=True,
                    processing_status='completed',
                    chunks_count=chunks_saved
                )
                
                caption_type = result.get('caption_type', 'unknown')
                caption_quality = 'HIGH (Manual Captions)' if caption_type == 'manual' else 'MEDIUM (Auto-Generated)'
                
                print(f"\n✅ YouTube Transcript Processing Complete!")
                print(f"   Method: YouTube API (no download)")
                print(f"   Title: {video_title}")
                print(f"   File ID: {file_obj.id}")
                print(f"   Chunks: {chunks_saved}")
                print(f"{'='*70}\n")
                
                return JsonResponse({
                    'success': True,
                    'method': 'transcript_api',
                    'file_id': file_obj.id,
                    'filename': file_obj.original_filename,
                    'title': video_title,
                    'chunks': chunks_saved,
                    'transcript_length': len(result['full_transcript']),
                    'video_id': video_id,
                    'caption_type': caption_type,
                    'caption_quality': caption_quality,
                    'message': f'✅ Processed using {caption_quality} - No download needed!'
                })
                
        except Exception as e:
            print(f"⚠️ YouTube Transcript API failed: {e}")
            print(f"   Falling back to download method...")
        
        print(f"\n📥 STRATEGY 2: Downloading video with yt-dlp...")
        
        try:
            import yt_dlp
        except ImportError:
            return JsonResponse({
                'success': False,
                'error': 'yt-dlp not installed. Please run: pip install yt-dlp'
            }, status=500)
        
        downloads_dir = settings.MEDIA_ROOT / 'uploads'
        downloads_dir.mkdir(parents=True, exist_ok=True)
        
        print("📋 Fetching video info...")
        with yt_dlp.YoutubeDL({'quiet': True}) as ydl:
            info = ydl.extract_info(url, download=False)
            video_title = info.get('title', f'YouTube_{video_id}')
        
        clean_title = re.sub(r'[^\w\s-]', '', video_title).strip()
        words = clean_title.split()[:4]  
        clean_title = '_'.join(words)
        
        if not clean_title:
            clean_title = f"youtube_{video_id}"
        
        print(f"   Video title: {video_title}")
        print(f"   Clean filename: {clean_title}")
        
        output_template = str(downloads_dir / f'{clean_title}.%(ext)s')
        
        def progress_hook(d):
            if d['status'] == 'downloading':
                percent = d.get('_percent_str', 'N/A')
                speed = d.get('_speed_str', 'N/A')
                print(f"   Downloading: {percent} at {speed}")
            elif d['status'] == 'finished':
                print(f"   ✅ Download complete, merging files...")
        
        ydl_opts = {
            'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best',
            'outtmpl': output_template,
            'noplaylist': True,
            'quiet': False,
            'no_warnings': False,
            'socket_timeout': 30,
            'max_filesize': 5 * 1024 * 1024 * 1024,
            'progress_hooks': [progress_hook],
            'keepvideo': False,
        }
        
        print("📥 Starting download...")
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            downloaded_file = ydl.prepare_filename(info)
        
        if not os.path.exists(downloaded_file):
            return JsonResponse({
                'success': False,
                'error': 'Video download failed - file not found'
            }, status=500)
        
        print(f"✅ Video downloaded successfully!")
        print(f"   Title: {video_title}")
        print(f"   File: {downloaded_file}")
        print(f"   Size: {os.path.getsize(downloaded_file) / (1024*1024):.2f} MB")
        
        from django.core.files import File
        
        filename = os.path.basename(downloaded_file)
        
        with open(downloaded_file, 'rb') as f:
            django_file = File(f, name=filename)
            file_obj = UploadedFile.objects.create(
                user=request.user,
                file_type='video',
                original_filename=f"{clean_title}.mp4",  
                processing_status='uploaded'
            )
            file_obj.file_path.save(filename, django_file, save=True)
        
        print(f"✅ Database record created (ID: {file_obj.id})")
        print(f"   Stored as: {clean_title}.mp4")
        print(f"{'='*70}\n")
        
        return JsonResponse({
            'success': True,
            'method': 'download',
            'file_id': file_obj.id,
            'filename': file_obj.original_filename,
            'title': video_title,
            'type': 'video',
            'message': '⚠️ Transcripts unavailable - Downloaded video for Whisper processing'
        })
        
    except Exception as e:
        print(f"❌ YouTube processing error: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'success': False,
            'error': f'Processing failed: {str(e)}'
        }, status=500)