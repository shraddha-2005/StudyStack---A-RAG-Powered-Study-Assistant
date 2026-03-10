from docx import Document
import re

def process_docx(file_path, filename):
    """Extract text from DOCX file"""
    try:
        print(f"📝 Opening DOCX file: {filename}")
        doc = Document(file_path)
        
        full_text = []
        paragraph_count = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
                paragraph_count += 1
        
        print(f"   - Extracted {paragraph_count} paragraphs")
        
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
                        table_count += 1
        
        print(f"   - Extracted {table_count} table cells")
        
        text = '\n\n'.join(full_text)
        
        text = re.sub(r'\n{3,}', '\n\n', text)  
        text = re.sub(r' {2,}', ' ', text)       
        text = text.strip()
        
        print(f"   - Total characters: {len(text)}")
        
        if len(text) < 50:
            return {
                'text': '',
                'title': filename.replace('.docx', '').replace('.doc', ''),
                'success': False,
                'error': 'Document appears to be empty or contains very little text'
            }
        
        return {
            'text': text,
            'title': filename.replace('.docx', '').replace('.doc', ''),
            'success': True
        }
        
    except Exception as e:
        print(f"❌ Error processing DOCX: {e}")
        import traceback
        traceback.print_exc()
        return {
            'text': '',
            'title': filename,
            'success': False,
            'error': str(e)
        }